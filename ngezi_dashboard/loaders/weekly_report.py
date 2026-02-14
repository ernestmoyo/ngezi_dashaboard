"""
Loaders for the weekly Excel report and Word narrative document.

Weekly Excel: October 2021 07.10.21.xlsx
    Sheets: Crushing and Milling, Metal Production, Consumables,
            FAZ Weekly report, Grind, DATA, DATA (2)

Weekly Word: October 2021 - Week 5r.docx
    Sections: Introduction, People Issues, Production, Projects table
"""

import logging
from datetime import datetime

import openpyxl
import pandas as pd

from .utils import safe_float, normalise_date

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Weekly Excel loaders
# ---------------------------------------------------------------------------

def load_weekly_excel(path: str) -> dict[str, pd.DataFrame]:
    """Load all relevant sheets from the weekly Excel report.

    Returns a dict mapping sheet category to DataFrame.
    """
    results = {}

    try:
        results["crushing_milling"] = load_weekly_crushing_milling(path)
    except Exception:
        logger.exception("Failed to load Crushing and Milling sheet")

    try:
        results["metal_production"] = load_weekly_metal_production(path)
    except Exception:
        logger.exception("Failed to load Metal Production sheet")

    try:
        results["consumables"] = load_weekly_consumables(path)
    except Exception:
        logger.exception("Failed to load Consumables sheet")

    try:
        results["grind_trend"] = load_grind_trend(path)
    except Exception:
        logger.exception("Failed to load Grind sheet")

    try:
        results["daily_data"] = load_daily_data(path)
    except Exception:
        logger.exception("Failed to load DATA (2) sheet")

    return results


def load_weekly_crushing_milling(path: str) -> pd.DataFrame:
    """Load Crushing and Milling sheet.

    Assumptions
    -----------
    - Crushing data: rows 6-11, with Week (cols C-F) and Month (cols G-J).
    - Milling data: rows 23-32, same column layout.
    - Row 5/22: column headers (Variable, Actual, Target, Var., % Var.)

    Returns
    -------
    DataFrame with columns:
        section, variable, week_actual, week_target, week_var, week_var_pct,
        month_actual, month_target, month_var, month_var_pct
    """
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb["Crushing and Milling"]

    rows = []

    # Crushing section: rows 6-11
    for row_idx in range(6, 12):
        var_name = ws.cell(row=row_idx, column=2).value
        if var_name is None:
            continue
        rows.append({
            "section": "crushing",
            "variable": str(var_name).strip(),
            "week_actual": safe_float(ws.cell(row=row_idx, column=3).value),
            "week_target": safe_float(ws.cell(row=row_idx, column=4).value),
            "week_var": safe_float(ws.cell(row=row_idx, column=5).value),
            "week_var_pct": safe_float(ws.cell(row=row_idx, column=6).value),
            "month_actual": safe_float(ws.cell(row=row_idx, column=7).value),
            "month_target": safe_float(ws.cell(row=row_idx, column=8).value),
            "month_var": safe_float(ws.cell(row=row_idx, column=9).value),
            "month_var_pct": safe_float(ws.cell(row=row_idx, column=10).value),
        })

    # Milling section: rows 23-32
    for row_idx in range(23, 33):
        var_name = ws.cell(row=row_idx, column=2).value
        if var_name is None:
            continue
        rows.append({
            "section": "milling",
            "variable": str(var_name).strip(),
            "week_actual": safe_float(ws.cell(row=row_idx, column=3).value),
            "week_target": safe_float(ws.cell(row=row_idx, column=4).value),
            "week_var": safe_float(ws.cell(row=row_idx, column=5).value),
            "week_var_pct": safe_float(ws.cell(row=row_idx, column=6).value),
            "month_actual": safe_float(ws.cell(row=row_idx, column=7).value),
            "month_target": safe_float(ws.cell(row=row_idx, column=8).value),
            "month_var": safe_float(ws.cell(row=row_idx, column=9).value),
            "month_var_pct": safe_float(ws.cell(row=row_idx, column=10).value),
        })

    wb.close()

    df = pd.DataFrame(rows)
    logger.info("Loaded %d crushing/milling rows from %s", len(df), path)
    return df


def load_weekly_metal_production(path: str) -> pd.DataFrame:
    """Load Metal Production sheet.

    Assumptions
    -----------
    - Metal assays rows 5-15 (elements + aggregate grades).
    - Production rows 17-22 (concentrate, ozs, recovery, Ni).
    - Filtration row 27 (cake moisture).
    - Columns: B=Variable, C=Actual, D=Target, E=Var, F=%Var.

    Returns
    -------
    DataFrame with columns: variable, actual, target, var, var_pct
    """
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb["Metal Production"]

    rows = []

    # Scan relevant rows
    for row_idx in list(range(5, 16)) + list(range(17, 23)) + [27]:
        var_name = ws.cell(row=row_idx, column=2).value
        if var_name is None:
            continue
        var_name = str(var_name).strip()
        if not var_name:
            continue

        rows.append({
            "variable": var_name,
            "actual": safe_float(ws.cell(row=row_idx, column=3).value),
            "target": safe_float(ws.cell(row=row_idx, column=4).value),
            "var": safe_float(ws.cell(row=row_idx, column=5).value),
            "var_pct": safe_float(ws.cell(row=row_idx, column=6).value),
        })

    wb.close()

    df = pd.DataFrame(rows)
    logger.info("Loaded %d metal production rows from %s", len(df), path)
    return df


def load_weekly_consumables(path: str) -> pd.DataFrame:
    """Load Consumables sheet (reagents + water consumption).

    Assumptions
    -----------
    - Reagent rows 5-15: B=name, C=Actual(g/t), D=Budget(g/t), E=Var, F=%Var
    - Water rows 25-27: B=name, C=Actual(m3/t), D=Budget(m3/t), E=Var, F=%Var

    Returns
    -------
    DataFrame with columns: category, consumable, actual, budget, var, var_pct
    """
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb["Consumables"]

    rows = []

    # Reagents: rows 5-15
    for row_idx in range(5, 16):
        name = ws.cell(row=row_idx, column=2).value
        if name is None:
            continue
        name = str(name).strip()
        if not name:
            continue

        rows.append({
            "category": "reagent",
            "consumable": name,
            "actual": safe_float(ws.cell(row=row_idx, column=3).value),
            "budget": safe_float(ws.cell(row=row_idx, column=4).value),
            "var": safe_float(ws.cell(row=row_idx, column=5).value),
            "var_pct": safe_float(ws.cell(row=row_idx, column=6).value),
        })

    # Water consumption: rows 25-27
    for row_idx in range(25, 28):
        name = ws.cell(row=row_idx, column=2).value
        if name is None:
            continue
        name = str(name).strip()
        if not name:
            continue

        rows.append({
            "category": "water",
            "consumable": name,
            "actual": safe_float(ws.cell(row=row_idx, column=3).value),
            "budget": safe_float(ws.cell(row=row_idx, column=4).value),
            "var": safe_float(ws.cell(row=row_idx, column=5).value),
            "var_pct": safe_float(ws.cell(row=row_idx, column=6).value),
        })

    wb.close()

    df = pd.DataFrame(rows)
    logger.info("Loaded %d consumable rows from %s", len(df), path)
    return df


def load_grind_trend(path: str) -> pd.DataFrame:
    """Load Grind sheet — weekly grind % and milling rate, plus monthly mill ball data.

    Assumptions
    -----------
    - Row 2: month names spanning cols F-AX (5 cols per month, Feb-Oct 2021).
    - Row 3: week labels (Week 1-5).
    - Row 4: grind values (%-75 microns).
    - Row 5: milling rate values (tph).
    - Rows 31-40: monthly mill ball consumption (actual vs budget=540 g/t).

    Returns
    -------
    DataFrame with columns: month, week, grind_pct, milling_rate_tph
    """
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb["Grind"]

    rows = []

    # Weekly grind/rate data: cols F(6) through AX(50)
    current_month = None
    for col_idx in range(6, 51):
        # Row 2: month name (only in first col of each group due to merge)
        month_val = ws.cell(row=2, column=col_idx).value
        if month_val is not None:
            current_month = str(month_val).strip()

        week_val = ws.cell(row=3, column=col_idx).value
        grind_val = safe_float(ws.cell(row=4, column=col_idx).value)
        rate_val = safe_float(ws.cell(row=5, column=col_idx).value)

        if grind_val is not None or rate_val is not None:
            rows.append({
                "month": current_month,
                "week": str(week_val).strip() if week_val else None,
                "grind_pct": grind_val,
                "milling_rate_tph": rate_val,
            })

    wb.close()

    df = pd.DataFrame(rows)
    logger.info("Loaded %d grind trend rows from %s", len(df), path)
    return df


def load_daily_data(path: str) -> pd.DataFrame:
    """Load DATA (2) sheet — daily production tracking for October 2021.

    Assumptions
    -----------
    - Rows 4-34: one row per day of October 2021.
    - Column D: date.
    - Column F: daily milled tonnage actual.
    - Column G: daily target.
    - Columns H-I: month-to-date actual and target.
    - Column J: month-to-date variance %.

    Returns
    -------
    DataFrame with columns:
        date, daily_actual, daily_target, mtd_actual, mtd_target, mtd_var_pct
    """
    wb = openpyxl.load_workbook(path, data_only=True)

    sheet_name = "DATA (2)"
    if sheet_name not in wb.sheetnames:
        logger.warning("Sheet '%s' not found. Available: %s", sheet_name, wb.sheetnames)
        wb.close()
        return pd.DataFrame()

    ws = wb[sheet_name]

    rows = []
    for row_idx in range(4, 35):
        date_val = ws.cell(row=row_idx, column=4).value
        if date_val is None:
            continue

        date_ts = normalise_date(date_val)

        daily_actual = safe_float(ws.cell(row=row_idx, column=6).value)
        daily_target = safe_float(ws.cell(row=row_idx, column=7).value)
        mtd_actual = safe_float(ws.cell(row=row_idx, column=8).value)
        mtd_target = safe_float(ws.cell(row=row_idx, column=9).value)
        mtd_var_pct = safe_float(ws.cell(row=row_idx, column=10).value)

        rows.append({
            "date": date_ts,
            "daily_actual": daily_actual,
            "daily_target": daily_target,
            "mtd_actual": mtd_actual,
            "mtd_target": mtd_target,
            "mtd_var_pct": mtd_var_pct,
        })

    wb.close()

    df = pd.DataFrame(rows)
    if not df.empty:
        df["date"] = pd.to_datetime(df["date"])

    logger.info("Loaded %d daily data rows from %s", len(df), path)
    return df


# ---------------------------------------------------------------------------
# Word document loader
# ---------------------------------------------------------------------------

def load_projects_from_docx(path: str) -> pd.DataFrame:
    """Extract the project tracker table from the weekly Word report.

    Assumptions
    -----------
    - The document contains exactly one table with 13 rows x 5 columns.
    - Columns: Item, Project, Responsibility, Completion date, Comments.
    - First row is the header; rows 2-13 are project items (a through l).

    Returns
    -------
    DataFrame with columns:
        project_id, project_name, responsible, planned_completion, status, comments
    """
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx is required to parse Word documents. pip install python-docx")
        raise

    try:
        doc = Document(path)
    except Exception:
        logger.exception("Failed to open Word document: %s", path)
        raise

    if not doc.tables:
        logger.warning("No tables found in %s", path)
        return pd.DataFrame()

    table = doc.tables[0]
    rows = []

    for i, row in enumerate(table.rows):
        if i == 0:
            # Skip header row
            continue

        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) < 5:
            continue

        item_letter = cells[0].lower().rstrip(".")
        project_name = cells[1]
        responsible = cells[2]
        completion_raw = cells[3]
        comments = cells[4]

        # Parse completion date
        planned_completion = _parse_completion_date(completion_raw)

        # Derive status from comments
        status = _derive_status(comments)

        rows.append({
            "project_id": item_letter,
            "project_name": project_name,
            "responsible": responsible,
            "planned_completion": planned_completion,
            "status": status,
            "comments": comments,
        })

    df = pd.DataFrame(rows)
    logger.info("Loaded %d project rows from %s", len(df), path)
    return df


def _parse_completion_date(raw: str) -> pd.Timestamp | None:
    """Try to parse a completion date string from the Word table."""
    if not raw:
        return None

    # Common formats in the source data:
    # "June 2021", "December 2021", "30 November 2021"
    # "08/06/21", "14/06/21" (dd/mm/yy)
    # "July 21", "Aug  21", "Sept 21", "Nov 21", "Dec 21"
    cleaned = " ".join(raw.strip().split())  # collapse multiple spaces

    formats = [
        "%B %Y",       # June 2021
        "%d %B %Y",    # 30 November 2021
        "%b %Y",       # Jun 2021
        "%d %b %Y",    # 30 Jun 2021
        "%Y-%m-%d",    # 2021-06-30
        "%d/%m/%y",    # 08/06/21
        "%d/%m/%Y",    # 08/06/2021
        "%B %y",       # June 21
        "%b %y",       # Jun 21
    ]

    # Handle abbreviated months with trailing year: "Sept 21", "Aug  21"
    month_aliases = {
        "Sept": "Sep",
        "Sept.": "Sep",
    }
    for alias, replacement in month_aliases.items():
        if alias in cleaned:
            cleaned = cleaned.replace(alias, replacement)

    for fmt in formats:
        try:
            return pd.Timestamp(datetime.strptime(cleaned, fmt))
        except ValueError:
            continue

    logger.warning("Could not parse completion date: '%s'", raw)
    return None


def _derive_status(comments: str) -> str:
    """Derive project status from the comments text."""
    if not comments:
        return "unknown"

    lower = comments.lower()
    if "complete" in lower or "done" in lower or "closed" in lower:
        return "completed"
    if "in progress" in lower or "ongoing" in lower or "underway" in lower:
        return "in_progress"
    if "pending" in lower or "awaiting" in lower or "not started" in lower:
        return "pending"
    if "delayed" in lower or "overdue" in lower:
        return "delayed"

    # Default: treat as in_progress if there are comments indicating activity
    if len(comments) > 5:
        return "in_progress"

    return "unknown"
